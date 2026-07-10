from __future__ import annotations

from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import or_

from backend.app.api import job as job_api
from backend.app.database import SessionLocal
from backend.app.main import app
from backend.app.models.ai_resume_analysis import AIResumeAnalysis
from backend.app.models.application import Application
from backend.app.models.candidate import Candidate
from backend.app.models.embedding import Embedding
from backend.app.models.job import Job
from backend.app.models.resume import Resume
from backend.app.models.semantic_similarity_result import SemanticSimilarityResult
from backend.app.models.user import User


async def fake_structure(*, resume_text: str):
    return (
        {
            "version": 1,
            "sections": {
                "skills": {"items": ["Python", "FastAPI", "MySQL"]},
                "experience": {"items": []},
                "projects": {"items": []},
                "education": {"items": []},
            },
        },
        {"model": "module-smoke"},
    )


async def fake_analysis(**kwargs):
    return (
        {
            "candidate_summary": "The candidate matches Python and FastAPI, but is missing no required skills.",
            "strengths": ["Backend API development"],
            "weaknesses": [],
            "matched_skills": kwargs.get("matched_skills", []),
            "missing_skills": kwargs.get("missing_skills", []),
            "recommendation": "Good Fit",
            "reasoning": "The candidate demonstrates relevant backend skills.",
        },
        {"status": "success", "model": "module-smoke"},
    )


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Module Verification Candidate", level=1)
    document.add_paragraph("Python FastAPI MySQL developer with API and database experience.")
    document.add_heading("Skills", level=2)
    document.add_paragraph("Python, FastAPI, MySQL, React")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def cleanup_previous_smoke_data() -> None:
    db = SessionLocal()
    try:
        users = db.query(User).filter(User.email.like("module-%@example.com")).all()
        candidates = db.query(Candidate).filter(Candidate.email.like("module-%@example.com")).all()
        jobs = db.query(Job).filter(Job.job_title == "Module Verification Engineer").all()
        user_ids = [int(row.id) for row in users]
        candidate_ids = [int(row.id) for row in candidates]
        job_ids = [int(row.id) for row in jobs]
        applications = (
            db.query(Application)
            .filter(or_(Application.job_id.in_(job_ids), Application.candidate_id.in_(candidate_ids)))
            .all()
        )
        resume_ids = [int(row.resume_id) for row in applications if row.resume_id]
        resumes = db.query(Resume).filter(
            or_(Resume.id.in_(resume_ids), Resume.candidate_id.in_(candidate_ids))
        ).all()
        resume_ids = [int(row.id) for row in resumes]
        for resume in resumes:
            path = Path(job_api.UPLOAD_DIR) / Path(resume.file_path)
            if path.exists():
                try:
                    path.unlink()
                except OSError:
                    pass
        if resume_ids:
            db.query(Embedding).filter(
                Embedding.entity_type == "resume",
                Embedding.entity_id.in_(resume_ids),
            ).delete(synchronize_session=False)
            db.query(SemanticSimilarityResult).filter(
                SemanticSimilarityResult.resume_id.in_(resume_ids)
            ).delete(synchronize_session=False)
        if job_ids:
            db.query(Embedding).filter(
                Embedding.entity_type == "job",
                Embedding.entity_id.in_(job_ids),
            ).delete(synchronize_session=False)
            db.query(SemanticSimilarityResult).filter(
                SemanticSimilarityResult.job_id.in_(job_ids)
            ).delete(synchronize_session=False)
        for application in applications:
            db.query(AIResumeAnalysis).filter(AIResumeAnalysis.application_id == application.id).delete(synchronize_session=False)
            db.delete(application)
        for resume in resumes:
            db.delete(resume)
        for job in jobs:
            db.delete(job)
        for candidate in candidates:
            db.delete(candidate)
        for user in users:
            db.delete(user)
        db.commit()
    except Exception:
        db.rollback()
        raise
    finally:
        db.close()


def main() -> None:
    cleanup_previous_smoke_data()
    suffix = uuid4().hex[:12]
    recruiter_email = f"module-recruiter-{suffix}@example.com"
    candidate_email = f"module-candidate-{suffix}@example.com"
    password = "ModuleCheck123!"
    created_job_id: int | None = None
    created_application_id: int | None = None

    job_api.ai_structure_resume = fake_structure
    job_api.analyze_resume_for_job = fake_analysis
    job_api.resume_job_similarity = lambda *_args, **_kwargs: 0.74
    job_api.get_or_create_embedding = lambda *_args, **_kwargs: None

    try:
        with TestClient(app) as client:
            recruiter_signup = client.post(
                "/auth/signup",
                json={"email": recruiter_email, "name": "Module Recruiter", "password": password, "role": "recruiter"},
            )
            assert recruiter_signup.status_code == 200, recruiter_signup.text

            candidate_signup = client.post(
                "/auth/signup",
                json={"email": candidate_email, "name": "Module Candidate", "password": password, "role": "candidate"},
            )
            assert candidate_signup.status_code == 200, candidate_signup.text

            recruiter_login = client.post(
                "/auth/login",
                json={"email": recruiter_email, "password": password, "role": "recruiter"},
            )
            candidate_login = client.post(
                "/auth/login",
                json={"email": candidate_email, "password": password, "role": "candidate"},
            )
            assert recruiter_login.status_code == 200, recruiter_login.text
            assert candidate_login.status_code == 200, candidate_login.text

            recruiter_token = recruiter_login.json()["access_token"]
            candidate_token = candidate_login.json()["access_token"]
            recruiter_headers = {"Authorization": f"Bearer {recruiter_token}"}
            candidate_headers = {"Authorization": f"Bearer {candidate_token}"}

            forbidden = client.post(
                "/jobs",
                headers=candidate_headers,
                json={"title": "Forbidden Job", "description": "Candidates cannot create this job."},
            )
            assert forbidden.status_code == 403, forbidden.text

            create_job = client.post(
                "/jobs",
                headers=recruiter_headers,
                json={
                    "title": "Module Verification Engineer",
                    "short_description": "Verifies Modules 1 through 5",
                    "description": "Build and maintain Python APIs backed by MySQL databases.",
                    "required_skills": ["Python", "FastAPI", "MySQL"],
                    "status": "active",
                },
            )
            assert create_job.status_code == 201, create_job.text
            job_payload = create_job.json()["job"]
            created_job_id = int(job_payload["id"])
            assert job_payload["required_skills"] == ["Python", "FastAPI", "MySQL"]

            standalone_upload = client.post(
                "/resumes/upload",
                headers=candidate_headers,
                files={"file": ("profile.docx", make_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert standalone_upload.status_code == 404, standalone_upload.text

            apply_response = client.post(
                f"/jobs/{created_job_id}/apply",
                headers=candidate_headers,
                files={"file": ("resume.docx", make_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert apply_response.status_code == 200, apply_response.text
            application = apply_response.json()["application"]
            created_application_id = int(application["id"])
            assert int(application["job_id"]) == created_job_id
            assert int(application["resume_id"]) > 0
            assert 0 <= int(application["final_score"]) <= 100

            duplicate = client.post(
                f"/jobs/{created_job_id}/apply",
                headers=candidate_headers,
                files={"file": ("resume-again.docx", make_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert duplicate.status_code == 200, duplicate.text
            assert duplicate.json().get("already_applied") is True

            details = client.get(
                f"/jobs/applications/{created_application_id}",
                headers=candidate_headers,
            )
            assert details.status_code == 200, details.text
            detail_payload = details.json()["application"]
            assert int(detail_payload["job_id"]) == created_job_id
            assert int(detail_payload["resume_id"]) == int(application["resume_id"])
            assert detail_payload["final_score"] == application["final_score"]
            assert detail_payload["ai_analysis"]["recommendation"] == "Good Fit"

            status_update = client.patch(
                f"/jobs/applications/{created_application_id}/status",
                headers=recruiter_headers,
                json={"status": "shortlisted"},
            )
            assert status_update.status_code == 200, status_update.text
            assert status_update.json()["application"]["status"] == "shortlisted"

            updated_details = client.get(
                f"/jobs/applications/{created_application_id}",
                headers=candidate_headers,
            )
            assert updated_details.status_code == 200, updated_details.text
            assert updated_details.json()["application"]["status"] == "shortlisted"
            assert updated_details.json()["application"]["final_score"] == application["final_score"]

            removed = client.delete(
                f"/jobs/applications/{created_application_id}",
                headers=candidate_headers,
            )
            assert removed.status_code == 200, removed.text
            created_application_id = None

        print("modules_1_5_smoke_ok")
    finally:
        db = SessionLocal()
        try:
            users = db.query(User).filter(User.email.in_([recruiter_email, candidate_email])).all()
            user_ids = [int(user.id) for user in users]
            candidates = db.query(Candidate).filter(Candidate.email == candidate_email).all()
            candidate_ids = [int(candidate.id) for candidate in candidates]

            if created_application_id:
                application = db.query(Application).filter(Application.id == created_application_id).first()
                if application:
                    db.query(AIResumeAnalysis).filter(AIResumeAnalysis.application_id == application.id).delete()
                    resume = db.query(Resume).filter(Resume.id == application.resume_id).first()
                    if resume:
                        path = Path(job_api.UPLOAD_DIR) / Path(resume.file_path)
                        if path.exists():
                            try:
                                path.unlink()
                            except OSError:
                                pass
                        db.delete(resume)
                    db.delete(application)

            if created_job_id:
                db.query(SemanticSimilarityResult).filter(SemanticSimilarityResult.job_id == created_job_id).delete()
                db.query(Embedding).filter(Embedding.entity_type == "job", Embedding.entity_id == created_job_id).delete()
                job = db.query(Job).filter(Job.id == created_job_id).first()
                if job:
                    db.delete(job)

            for candidate_id in candidate_ids:
                resumes = db.query(Resume).filter(Resume.candidate_id == candidate_id).all()
                for resume in resumes:
                    path = Path(job_api.UPLOAD_DIR) / Path(resume.file_path)
                    if path.exists():
                        try:
                            path.unlink()
                        except OSError:
                            pass
                    db.query(Embedding).filter(
                        Embedding.entity_type == "resume",
                        Embedding.entity_id == int(resume.id),
                    ).delete()
                    db.delete(resume)
                candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
                if candidate:
                    db.delete(candidate)
            for user_id in user_ids:
                user = db.query(User).filter(User.id == user_id).first()
                if user:
                    db.delete(user)
            db.commit()
        except Exception:
            db.rollback()
            raise
        finally:
            db.close()


if __name__ == "__main__":
    main()
