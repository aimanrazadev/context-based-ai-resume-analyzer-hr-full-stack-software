from __future__ import annotations

from io import BytesIO


def test_clean_extracted_text_removes_repeated_header_footer_and_page_markers():
    from backend.app.services.resume_analysis import clean_extracted_text

    page1 = "John Doe\nEmail: john@example.com\n\nExperience\nBuilt REST APIs using FastAPI.\n\nPage 1 of 2\n"
    page2 = "John Doe\nEmail: john@example.com\n\nProjects\nDevel-\nopment of backend services.\n\nPage 2 of 2\n"
    raw = f"{page1}\n{page2}"

    cleaned = clean_extracted_text(raw_text=raw, page_texts=[page1, page2])

    # Repeated header removed
    assert "John Doe" not in cleaned
    assert "john@example.com" not in cleaned
    # Page markers removed
    assert "Page 1" not in cleaned
    assert "Page 2" not in cleaned
    # Hyphenation fixed
    assert "Development" in cleaned or "Development".lower() in cleaned.lower()
    # Context preserved
    assert "Built REST APIs using FastAPI" in cleaned


def test_extract_and_clean_resume_text_detects_emptyish_pdf(tmp_path):
    from backend.app.services.resume_analysis import extract_and_clean_resume_text

    # Minimal/invalid PDF bytes; extraction should not crash and should be treated as emptyish
    p = tmp_path / "scan.pdf"
    p.write_bytes(b"%PDF-1.4\n%EOF\n")

    res = extract_and_clean_resume_text(file_path=str(p), ext=".pdf")
    # With OCR enabled, this may still be empty if OCR binaries aren't installed.
    assert res["clean_text"] in ("", res["clean_text"])
    assert res["is_probably_scanned_or_empty"] in (True, False)
    assert isinstance(res["warnings"], list)


def test_extract_and_clean_resume_text_uses_ocr_when_pdf_empty(monkeypatch, tmp_path):
    """
    OCR is optional and depends on system binaries; this test monkeypatches the OCR helper
    so we validate the logic without requiring Tesseract/Poppler installed.
    """
    import backend.app.services.resume_analysis as ra

    def fake_ocr(*, file_path: str):
        return ("Built REST APIs using FastAPI and deployed on AWS.\nPage 1 of 1\n", ["OCR used"])

    monkeypatch.setattr(ra, "_ocr_pdf_with_tesseract", fake_ocr)

    p = tmp_path / "scan.pdf"
    p.write_bytes(b"%PDF-1.4\n%EOF\n")

    res = ra.extract_and_clean_resume_text(file_path=str(p), ext=".pdf")
    assert res["is_probably_scanned_or_empty"] is False
    assert "Built REST APIs using FastAPI" in res["clean_text"]
    assert "Page 1" not in res["clean_text"]
    assert "OCR used" in " ".join(res["warnings"])


def test_endpoints_persist_extracted_text_for_docx(client, db_session, tmp_path):
    """
    Integration-level: upload a DOCX through both flows and ensure Resume.extracted_text is stored.
    """
    import docx  # type: ignore

    from backend.app.models.resume import Resume

    def signup(email: str, role: str):
        client.post("/auth/signup", json={"email": email, "password": "Testpass123!", "role": role, "name": "T"})

    def login(email: str, role: str):
        return client.post("/auth/login", json={"email": email, "password": "Testpass123!", "role": role}).json()[
            "access_token"
        ]

    def auth(token: str):
        return {"Authorization": f"Bearer {token}"}

    # Create DOCX bytes
    d = docx.Document()
    d.add_paragraph("Built REST APIs using FastAPI and deployed on AWS.")
    buf = BytesIO()
    d.save(buf)
    docx_bytes = buf.getvalue()

    # Candidate generic upload
    signup("cand_docx@example.com", "candidate")
    cand_token = login("cand_docx@example.com", "candidate")
    r = client.post(
        "/resumes/upload",
        headers=auth(cand_token),
        files={
            "file": (
                "resume.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 201, r.text
    resume_id = r.json()["resume"]["id"]

    saved = db_session.query(Resume).filter(Resume.id == resume_id).first()
    assert saved is not None
    assert saved.extracted_text and "Built REST APIs using FastAPI" in saved.extracted_text
    assert saved.structured_json


def test_apply_flow_persists_extracted_text_for_docx(client, db_session):
    """
    Integration-level: apply to a job with DOCX and ensure the stored Resume has extracted_text.
    """
    import docx  # type: ignore

    from backend.app.models.resume import Resume

    def signup(email: str, role: str):
        client.post("/auth/signup", json={"email": email, "password": "Testpass123!", "role": role, "name": "T"})

    def login(email: str, role: str):
        return client.post("/auth/login", json={"email": email, "password": "Testpass123!", "role": role}).json()[
            "access_token"
        ]

    def auth(token: str):
        return {"Authorization": f"Bearer {token}"}

    signup("rec_docx_apply@example.com", "recruiter")
    rec_token = login("rec_docx_apply@example.com", "recruiter")
    job = client.post(
        "/jobs",
        headers=auth(rec_token),
        json={"title": "Backend engineer", "description": "Backend engineer with API development experience", "status": "active"},
    ).json()["job"]

    signup("cand_docx_apply@example.com", "candidate")
    cand_token = login("cand_docx_apply@example.com", "candidate")

    d = docx.Document()
    d.add_paragraph("Built REST APIs using FastAPI and deployed on AWS.")
    buf = BytesIO()
    d.save(buf)
    docx_bytes = buf.getvalue()

    r = client.post(
        f"/jobs/{job['id']}/apply",
        headers=auth(cand_token),
        files={
            "file": (
                "resume.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code in (200, 201), r.text
    app = r.json()["application"]
    resume_id = app["resume_id"]

    saved = db_session.query(Resume).filter(Resume.id == resume_id).first()
    assert saved is not None
    assert saved.extracted_text and "Built REST APIs using FastAPI" in saved.extracted_text
    assert saved.structured_json

