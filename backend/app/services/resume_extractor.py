"""
Resume extraction and cleaning utilities.

This module is responsible for:
- extracting text from uploaded PDF and DOCX resumes
- cleaning and normalizing noisy extraction output
- detecting scanned PDFs and rejecting them with actionable diagnostics
- returning extraction diagnostics for downstream parsing and AI stages
"""

import re


_WORD_RE = re.compile(r"[a-zA-Z0-9+#.]{2,}")

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_MULTISPACE_RE = re.compile(r"[ \t]{2,}")
_MULTINEWLINE_RE = re.compile(r"\n{3,}")
_BULLET_RE = re.compile(r"^[\s\u2022\u00b7\u25cf\u25e6\u25aa\u25ab\u2219\u2043\u2023]+", re.MULTILINE)
_PAGE_MARKER_RE = re.compile(
    r"^\s*(page\s*\d+(\s*of\s*\d+)?)\s*$|^\s*\d+\s*/\s*\d+\s*$|^\s*[-\u2013\u2014]{0,3}\s*\d+\s*[-\u2013\u2014]{0,3}\s*$",
    re.IGNORECASE | re.MULTILINE,
)
_HYPHEN_LINEBREAK_RE = re.compile(r"([A-Za-z])-\n([A-Za-z])")
_MIN_NONWS_CHARS = 30
_ARTIFACT_REPLACEMENTS = {
    "\u2018": "'",
    "\u2019": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2013": "-",
    "\u2014": "-",
    "\u2026": "...",
    "\u00a0": " ",
    "\u00e2\u20ac\u201c": "-",
    "\u00e2\u20ac\u201d": "-",
    "\u00e2\u20ac\u0153": '"',
    "\u00e2\u20ac\u009d": '"',
    "\u00e2\u20ac\u02dc": "'",
    "\u00e2\u20ac\u2122": "'",
    "\u00e2\u20ac\u00a6": "...",
    "\u00e2\u20ac\u00a2": "-",
}


def extract_text_from_file(*, file_path: str, ext: str | None) -> str:
    """
    Extract text from a supported resume format.

    Extraction order:
    - PDF: use `PyMuPDF`
    - DOCX: use `python-docx`
    Unsupported or corrupt files return an empty string.
    """
    ext_norm = (ext or "").lower()

    if ext_norm == ".pdf":
        try:
            return "\n".join(extract_text_from_pdf_pages(file_path=file_path)).strip()
        except Exception:
            pass

    if ext_norm == ".docx":
        try:
            return extract_text_from_docx(file_path=file_path)
        except Exception:
            pass

    return ""


def extract_text_from_pdf_pages(*, file_path: str) -> list[str]:
    """
    Extract per-page text from a PDF using `PyMuPDF`.

    Args:
        file_path: Path to the uploaded PDF file.

    Returns:
        A list containing one extracted text string per page.

    Side Effects:
        Reads the PDF from disk.

    Error Handling:
        Raises import or file parsing errors to the caller, but converts
        per-page extraction failures into empty strings.
    """
    import fitz  # type: ignore

    out: list[str] = []
    with fitz.open(file_path) as doc:
        for page in doc:
            try:
                out.append(page.get_text("text") or "")
            except Exception:
                out.append("")
    return out


def extract_text_from_docx(*, file_path: str) -> str:
    """
    Extract visible paragraph text from a DOCX file.
    """
    import docx  # type: ignore

    d = docx.Document(file_path)
    parts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            line = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def _normalize_newlines(text: str) -> str:
    """Normalize newline sequences to `\\n`."""
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _strip_control_chars(text: str) -> str:
    """Remove low-level control characters."""
    return _CONTROL_CHARS_RE.sub("", text)


def _normalize_punctuation_and_symbols(text: str) -> str:
    """
    Normalize common punctuation, symbol, and mojibake artifacts.
    """
    out = text
    for bad, good in _ARTIFACT_REPLACEMENTS.items():
        out = out.replace(bad, good)
    return out


def _normalize_whitespace(text: str) -> str:
    """Collapse redundant spaces and excessive blank lines."""
    text = _MULTISPACE_RE.sub(" ", text)
    text = _MULTINEWLINE_RE.sub("\n\n", text)
    return text.strip()


def _fix_hyphenation(text: str) -> str:
    """Join words split across line breaks."""
    return _HYPHEN_LINEBREAK_RE.sub(r"\1\2", text)


def _normalize_bullets(text: str) -> str:
    """Normalize leading bullet glyphs into a plain dash."""
    return _BULLET_RE.sub("- ", text)


def _remove_page_markers(text: str) -> str:
    """Remove standalone page number markers from extracted text."""
    return _PAGE_MARKER_RE.sub("", text)


def _repeated_header_footer_lines(page_texts: list[str]) -> set[str]:
    """
    Identify repeated top/bottom page lines that likely represent headers or footers.
    """
    if len(page_texts) < 2:
        return set()

    def norm_line(line: str) -> str:
        line = _normalize_newlines(line)
        line = _strip_control_chars(line)
        line = _normalize_punctuation_and_symbols(line)
        line = _MULTISPACE_RE.sub(" ", line).strip()
        return line

    counts: dict[str, int] = {}
    for p in page_texts:
        lines = [norm_line(l) for l in _normalize_newlines(p).split("\n") if norm_line(l)]
        if not lines:
            continue
        head = lines[:4]
        tail = lines[-4:] if len(lines) >= 4 else lines
        for l in set(head + tail):
            if len(l) < 6 or l.isdigit():
                continue
            counts[l] = counts.get(l, 0) + 1

    return {l for (l, c) in counts.items() if c >= 2}


def clean_extracted_text(*, raw_text: str, page_texts: list[str] | None = None) -> str:
    """
    Clean extracted resume text while preserving recruiter-relevant structure.
    """
    text = raw_text or ""
    text = _normalize_newlines(text)
    text = _strip_control_chars(text)
    text = _normalize_punctuation_and_symbols(text)

    if page_texts:
        remove = _repeated_header_footer_lines(page_texts)
        if remove:
            cleaned_lines = []
            for l in text.split("\n"):
                nl = _MULTISPACE_RE.sub(" ", l).strip()
                if nl and nl in remove:
                    continue
                cleaned_lines.append(l)
            text = "\n".join(cleaned_lines)

    text = _remove_page_markers(text)
    text = _fix_hyphenation(text)
    text = _normalize_bullets(text)
    text = _normalize_whitespace(text)
    return text


def extract_and_clean_resume_text(*, file_path: str, ext: str | None) -> dict:
    """
    Extract, clean, and summarize resume text in a single service call.
    """
    warnings: list[str] = []
    quality_flags: list[str] = []
    ext_norm = (ext or "").lower()

    raw_text = ""
    page_texts: list[str] | None = None
    extraction_method = "unknown"
    page_count = 0
    error_code: str | None = None
    error_message: str | None = None

    try:
        if ext_norm == ".pdf":
            try:
                page_texts = extract_text_from_pdf_pages(file_path=file_path)
                page_count = len(page_texts)
                if page_count != 1:
                    error_code = "invalid_page_count"
                    error_message = "Resume should be exactly 1 page."
                    raw_text = "\n\n".join(page_texts).strip()
                    extraction_method = "pymupdf_pdf_text"
                else:
                    raw_text = "\n\n".join(page_texts).strip()
                    extraction_method = "pymupdf_pdf_text"
            except Exception as exc:
                warnings.append("Failed to parse PDF text with PyMuPDF.")
                error_code = "corrupt_file"
                error_message = "Could not read resume. Please upload a valid PDF or DOCX."
                raw_text = ""
        elif ext_norm == ".docx":
            try:
                raw_text = extract_text_from_docx(file_path=file_path)
                extraction_method = "docx_text"
            except Exception as exc:
                warnings.append("Failed to parse DOCX text with python-docx.")
                error_code = "corrupt_file"
                error_message = "Could not read resume. Please upload a valid PDF or DOCX."
                raw_text = ""
        else:
            raw_text = extract_text_from_file(file_path=file_path, ext=ext_norm)
            extraction_method = "generic_fallback"
    except Exception:
        warnings.append("Unexpected error during extraction.")
        raw_text = ""

    clean = clean_extracted_text(raw_text=raw_text, page_texts=page_texts)

    non_ws = re.sub(r"\s+", "", clean)
    is_emptyish = len(non_ws) < _MIN_NONWS_CHARS
    if is_emptyish:
        if ext_norm == ".pdf":
            warnings.append("PDF appears scanned or has no extractable text.")
            error_code = error_code or "scanned_pdf"
            error_message = error_message or "This resume appears to be scanned. Please upload a text-based PDF or DOCX."
        elif ext_norm == ".docx":
            warnings.append("DOCX appears empty or has no extractable text.")
        else:
            warnings.append("File has no extractable text.")
        if is_emptyish:
            clean = ""

    word_count = len(_WORD_RE.findall(clean))
    char_count = len(clean)
    alpha_chars = sum(1 for ch in clean if ch.isalpha())
    digit_chars = sum(1 for ch in clean if ch.isdigit())
    alpha_ratio = (alpha_chars / max(1, char_count)) if char_count else 0.0
    digit_ratio = (digit_chars / max(1, char_count)) if char_count else 0.0

    if clean and word_count < 25:
        quality_flags.append("very_low_word_count")
    if clean and alpha_ratio < 0.45:
        quality_flags.append("low_alpha_ratio")
    if clean and digit_ratio > 0.30:
        quality_flags.append("high_digit_ratio")
    if clean and len(clean.splitlines()) <= 2 and word_count < 40:
        quality_flags.append("weak_document_structure")
    if is_emptyish:
        quality_flags.append("empty_or_scanned")

    is_low_confidence = bool(
        is_emptyish
        or "very_low_word_count" in quality_flags
        or "low_alpha_ratio" in quality_flags
    )
    extraction_status = "success"
    if not clean or error_code:
        extraction_status = "failed"
    elif is_low_confidence:
        extraction_status = "low_confidence"

    return {
        "raw_text": raw_text,
        "clean_text": clean,
        "raw_len": len(raw_text or ""),
        "clean_len": len(clean or ""),
        "word_count": word_count,
        "char_count": char_count,
        "extraction_method": extraction_method,
        "page_count": page_count,
        "is_scanned": error_code == "scanned_pdf",
        "is_probably_scanned_or_empty": bool(is_emptyish),
        "is_low_confidence": is_low_confidence,
        "extraction_status": extraction_status,
        "quality_flags": quality_flags,
        "warnings": warnings,
        "error_code": error_code,
        "error_message": error_message,
    }
